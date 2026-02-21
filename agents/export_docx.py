from __future__ import annotations
import os, re
from docx import Document

def markdown_to_docx(report_md_path: str, out_docx_path: str) -> str:
    with open(report_md_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    sup_re = re.compile(r"<sup>(\d+)</sup>")

    for ln in lines:
        if ln.startswith("# "):
            doc.add_heading(ln[2:], level=1)
        elif ln.startswith("## "):
            doc.add_heading(ln[3:], level=2)
        elif re.match(r"^\d+\.\s", ln):
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_superscript(p, ln)
        elif ln.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_superscript(p, ln[2:])
        elif ln.strip() == "":
            doc.add_paragraph("")
        else:
            p = doc.add_paragraph()
            _add_runs_with_superscript(p, ln)

    os.makedirs(os.path.dirname(out_docx_path), exist_ok=True)
    doc.save(out_docx_path)
    return out_docx_path

def _add_runs_with_superscript(paragraph, text):
    parts = re.split(r"(<sup>\d+</sup>)", text)
    for part in parts:
        m = re.fullmatch(r"<sup>(\d+)</sup>", part)
        if m:
            r = paragraph.add_run(m.group(1))
            r.font.superscript = True
        else:
            paragraph.add_run(part)
